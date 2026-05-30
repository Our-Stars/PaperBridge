from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
import re

from docx.shared import Inches

from paperbridge.models import Document, Figure, Table

# XML 不允许的控制字符（除了 tab、LF、CR）
_XML_ILLEGAL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _safe(text: str) -> str:
    return _XML_ILLEGAL_RE.sub("", text)


def export_docx(document: Document, path: Path, base_dir: Path) -> None:
    docx = DocxDocument()
    figures_by_caption = {figure.caption_block_id: figure for figure in document.figures if figure.caption_block_id}
    tables_by_caption = {table.caption_block_id: table for table in document.tables if table.caption_block_id}
    emitted_figures: set[str] = set()
    emitted_tables: set[str] = set()

    for block in document.body_blocks:
        if block.id in figures_by_caption:
            _append_figure(docx, figures_by_caption[block.id], base_dir)
            emitted_figures.add(figures_by_caption[block.id].id)
            continue
        if block.id in tables_by_caption:
            _append_table(docx, tables_by_caption[block.id], base_dir)
            emitted_tables.add(tables_by_caption[block.id].id)
            continue

        if block.type == "title":
            docx.add_heading(_safe(document.metadata.title or block.text), level=0)
        elif block.type == "subtitle":
            docx.add_paragraph(_safe(block.text), style="Subtitle")
        elif block.type == "abstract_heading":
            docx.add_heading("Abstract", level=1)
        elif block.type == "heading_1":
            docx.add_heading(_safe(block.text), level=1)
        elif block.type == "heading_2":
            docx.add_heading(_safe(block.text), level=2)
        elif block.type == "heading_3":
            docx.add_heading(_safe(block.text), level=3)
        elif block.type == "reference_heading":
            docx.add_heading("References", level=1)
        elif block.type == "list_item":
            docx.add_paragraph(_safe(block.text.lstrip("-*• ")), style="List Bullet")
        elif block.type in {"paragraph", "abstract", "reference_item", "equation", "unknown"}:
            docx.add_paragraph(_safe(block.text))

    for figure in document.figures:
        if figure.id not in emitted_figures:
            _append_figure(docx, figure, base_dir)
    for table in document.tables:
        if table.id not in emitted_tables:
            _append_table(docx, table, base_dir)

    path.parent.mkdir(parents=True, exist_ok=True)
    docx.save(path)


def _append_figure(docx: DocxDocument, figure: Figure, base_dir: Path) -> None:
    if figure.image_path:
        image_path = base_dir / figure.image_path
        if image_path.exists():
            docx.add_picture(str(image_path), width=Inches(5.8))
    if figure.caption:
        docx.add_paragraph(_safe(figure.caption), style="Caption")


def _append_table(docx: DocxDocument, table: Table, base_dir: Path) -> None:
    if table.representation == "structured" and table.columns:
        if table.caption:
            docx.add_paragraph(_safe(table.caption), style="Caption")
        docx_table = docx.add_table(rows=1, cols=len(table.columns))
        docx_table.style = "Table Grid"
        for index, column in enumerate(table.columns):
            docx_table.rows[0].cells[index].text = _safe(column)
        for row in table.rows:
            cells = docx_table.add_row().cells
            for index, value in enumerate(row[: len(cells)]):
                cells[index].text = _safe(value)
    elif table.image_path:
        image_path = base_dir / table.image_path
        if image_path.exists():
            docx.add_picture(str(image_path), width=Inches(5.8))
        if table.caption:
            docx.add_paragraph(_safe(table.caption), style="Caption")
    elif table.caption:
        docx.add_paragraph(_safe(table.caption), style="Caption")
